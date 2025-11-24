
import streamlit as st
import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
import pdfkit
import tempfile
import os

st.set_page_config(page_title="Webpage to PDF & Word Extractor", layout="centered")
st.title("Webpage to PDF & Word Extractor")

st.markdown("""
This tool extracts the content of a public webpage as:
- A PDF that looks like the actual webpage (with images, layout, and CSS)
- A Word document with structured content (headings, lists, paragraphs)

**No code is visible to users. No installation is needed for your team.**
""")

url = st.text_input("Enter the webpage URL (include https://):")

def add_html_to_docx(soup, doc):
    # Walk through the HTML and add structure to the Word doc
    for element in soup.body.descendants:
        if isinstance(element, Tag):
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(strip=True), level=level)
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style='List Number')

def generate_docx(soup, url):
    doc = Document()
    doc.add_heading(f'Website Content: {url}', 0)
    try:
        add_html_to_docx(soup, doc)
    except Exception as e:
        doc.add_paragraph("Could not structure content. Raw text below:")
        doc.add_paragraph(soup.get_text(separator='\n'))
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    return temp_docx.name

def generate_pdf(url):
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    try:
        # Use default config; Streamlit Cloud has wkhtmltopdf pre-installed
        pdfkit.from_url(url, temp_pdf.name)
        return temp_pdf.name
    except Exception as e:
        return None

if st.button("Extract and Download"):
    if not url or not url.startswith("http"):
        st.error("Please enter a valid URL (including https://)")
    else:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        try:
            response = requests.get(url, headers=headers, timeout=20)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            docx_file = generate_docx(soup, url)
            pdf_file = generate_pdf(url)
            st.success("Extraction complete! Download your files below.")
            with open(docx_file, "rb") as f:
                st.download_button("Download Word Document", f, file_name="webpage_content.docx")
            if pdf_file and os.path.exists(pdf_file):
                with open(pdf_file, "rb") as f:
                    st.download_button("Download PDF (Webpage Look)", f, file_name="webpage_content.pdf")
            else:
                st.warning("PDF generation failed (site may block rendering or require login).")
        except requests.exceptions.RequestException as e:
            st.error(f"Network or access error: {e}")
        except Exception as e:
            st.error(f"Unexpected error: {e}")
